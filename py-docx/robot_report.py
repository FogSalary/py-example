import os
from docx import Document
import shutil
from datetime import datetime
from docx.shared import Inches, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_table_no_borders(table):
    """
    设置表格为无边框
    """
    tbl = table._element  # 获取表格的底层 XML 元素
    tbl_pr = tbl.xpath('.//w:tblPr')[0]  # 获取表格属性部分 <w:tblPr>
    
    # 创建一个新的无边框 <w:tblBorders> 元素
    tbl_borders = OxmlElement('w:tblBorders')
    
    # 定义所有边框为无 (w:val="none")
    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        border = OxmlElement(f'w:{border_name}')
        border.set(qn('w:val'), 'none')  # 使用 qn 来设置命名空间的属性
        tbl_borders.append(border)
    
    # 替换或添加边框设置
    tbl_pr.append(tbl_borders)

def add_page1(target_paragraph, document):
    # add sub title
    # add description and illustration
    # add bar plot figure
    # add table show data information
    # new_paragraph = paragraph.insert_paragraph_before()
    
    # == input param
    global pic_square_path, pic_path
    description = "This is a demostration about circle cmd."
    
    new_paragraph = document.add_heading("Circle cmd overview", level=3)
    target_paragraph._element.addprevious(new_paragraph._element)

    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    cell = table.cell(0, 0)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(description)
    cell = table.cell(0, 1)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(pic_square_path, width=Inches(3))
    target_paragraph._element.addprevious(table._element)

    paragraph = document.add_paragraph("")
    target_paragraph._element.addprevious(paragraph._element)

    paragraph = document.add_paragraph()
    paragraph.add_run().add_picture(pic_path, width=Inches(6))
    target_paragraph._element.addprevious(paragraph._element)

    table = document.add_table(rows=2, cols=3)
    table.style = "Table Grid"
    target_paragraph._element.addprevious(table._element)

    brk_page = document.add_page_break()
    target_paragraph._element.addprevious(brk_page._element)

def add_page2(target_paragraph, document):
    # add sub title
    # add 1x2 table pic show: trajectory 3D view, top view
    # add trajectory error curve
    # add table show data information

    new_paragraph = document.add_heading("Circle cmd overview", level=3)
    target_paragraph._element.addprevious(new_paragraph._element)

    table_pic = document.add_table(1, 2)
    # table_pic.style = "Table Grid"
    cell = table_pic.cell(0, 0)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(pic_square_path, width=Inches(3))
    cell = table_pic.cell(0, 1)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(pic_square_path, width=Inches(3))
    # set_table_no_borders(table_pic)
    target_paragraph._element.addprevious(table_pic._element)

    paragraph = document.add_paragraph("")
    target_paragraph._element.addprevious(paragraph._element)
    

    paragraph = document.add_paragraph()
    paragraph.add_run().add_picture(pic_path, width=Inches(6))
    target_paragraph._element.addprevious(paragraph._element)

    table = document.add_table(rows=2, cols=3)
    table.style = "Table Grid"
    target_paragraph._element.addprevious(table._element)

    brk_page = document.add_page_break()
    target_paragraph._element.addprevious(brk_page._element)


now = datetime.now()
current_path = os.path.dirname(__file__)
template_path = os.path.join(current_path, 'template.docx')
date_str = now.strftime("%Y-%m-%d-%H-%M-%S")
report_path = os.path.join(current_path, "New_%s.docx" % date_str)
shutil.copy(template_path, report_path)

pic_path = os.path.join(current_path, 'fig.png')
pic_square_path = os.path.join(current_path, 'fig1.png')

document = Document(report_path)


for index, paragraph in enumerate(document.paragraphs):
    if paragraph.text == "$circle_big":
        add_page1(paragraph, document)

        add_page2(paragraph, document)

        # empty tag paragraph
        p = paragraph._element
        p.getparent().remove(p)
        p._element = None


    # if paragraph.text == "$circle_big":
    #     new_paragraph = paragraph.insert_paragraph_before()
    #     new_paragraph.add_run("hello world")
        
    #     new_paragraph = paragraph.insert_paragraph_before()
    #     new_paragraph.add_run().add_picture(pic_square_path, height=Inches(2))

    #     # new_paragraph.add_run()
    #     table = document.add_table(3, 2)
    #     paragraph._element.addprevious(table._element)
    
    # elif paragraph.text == "$small_big":
    #     new_paragraph = paragraph.insert_paragraph_before()
    #     new_paragraph.add_run("hello world")

    #     pic_paragraph = document.add_paragraph()
    #     pic_paragraph.add_run().add_picture(pic_path, height=Inches(1))
    #     new_paragraph._element.addprevious(pic_paragraph._element)

    #     table = document.add_table(rows=1, cols=2)
    #     new_paragraph._element.addnext(table._element)


document.save(report_path)
        