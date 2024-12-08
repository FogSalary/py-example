import os
from docx import Document


current_path = os.path.dirname(__file__)
document = Document()

document.add_heading("1 Big Circle CP Performance", level=1)
document.add_heading("1.1 Start from base.", level=1)
paragraph = document.add_paragraph("$circle_big")

document.add_page_break()


document.add_heading("2 Small Circle CP Performance", level=1)
paragraph = document.add_paragraph("$small_big")
document.add_page_break()


document.add_heading("3 Diag Line CP Performance", level=1)
paragraph = document.add_paragraph("$line")
document.add_page_break()


save_path = os.path.join(current_path, "template.docx")
document.save(save_path)