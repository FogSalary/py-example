import os
from docx import Document
from docx.shared import Inches


def add_description_demopic_in_docx(paragraph_index, document, description, pic_path):
    table = document.add_table(rows=1, cols=2)
    cell = table.cell(0, 0)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(description)
    cell = table.cell(0, 1)
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run()
    run.add_picture(pic_path, width=Inches(2))

    document.element.body.insert(paragraph_index+1, table._tbl)


current_path = os.path.dirname(__file__)

document = Document()  # create a new/empty docx file
# document = Document(docx_path)  # open a docx file
document.add_heading('Document Title', 0)

p = document.add_paragraph('A plain paragraph having some ')  # adding a paragraph
p.add_run('bold').bold = True
p.add_run(' and some ')
p.add_run('italic.').italic = True

prior_paragraph = p.insert_paragraph_before("Insert one line before p paragraph")

document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph('first item in unordered', style='List Bullet')
document.add_paragraph('first item in ordered list', style='List Number')

# document.add_picture(os.path.join(current_path, 'fig.png'))
pic_path = os.path.join(current_path, 'fig.png')
pic_square_path = os.path.join(current_path, 'fig1.png')
document.add_picture(pic_path, width=Inches(6))
document.add_picture(pic_path, height=Inches(2.5))

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

document.add_page_break()  # adding a page break

paragraph = document.add_paragraph('Normal text, ')
paragraph.add_run('text with emphasis.', 'Emphasis')


insert_index = len(document.paragraphs)
paragraph_insert = document.add_paragraph('Normal text, ')
run = paragraph_insert.add_run('text with emphasis.')
run.style = 'Emphasis'
print(len(document.paragraphs))

document.add_page_break()
for i in range(10):
    document.add_heading("This is level %d heading" % i, level=i)


# =====================================
# ==========    Table     =============
# =====================================
document.add_page_break()
document.add_heading("Table Demostration", level=0)

table2 = document.add_table(rows=2, cols=2)
table2.style = 'Table Grid'

# acess table cell by one
cell = table2.cell(0, 1)
cell.text = "parrot, possibly dead"

row = table2.rows[1]
row.cells[0].text = 'Foo bar to you'
row.cells[1].text = 'And a hearty foo bar to you too sir!'

table3 = document.add_table(rows=2, cols=2)
col = table3.columns[0]
col.cells[0].text = 'First'
col.cells[1].text = 'Second'

# table4 = document.add_table(rows=2, cols=3)
for row in table.rows:
    for cell in row.cells:
        print(cell.text)

row_count = len(table.rows)
col_count = len(table.columns)


# get table data -------------
items = (
    (7, '1024', 'Plush kittens'),
    (3, '2042', 'Furbees'),
    (1, '1288', 'French Poodle Collars, Deluxe'),
)

# add table ------------------
table4 = document.add_table(1, 3)

# populate header row --------
heading_cells = table4.rows[0].cells
heading_cells[0].text = 'Qty'
heading_cells[1].text = 'SKU'
heading_cells[2].text = 'Description'

# add a data row for each item
for item in items:
    cells = table4.add_row().cells
    cells[0].text = str(item[0])
    cells[1].text = item[1]
    cells[2].text = item[2]

# 列出所有的表格样式
# table_styles = [style.name for style in document.styles]  # 6 是表格样式
# # table_styles = [style.name for style in document.styles if style.type == 6]  # 6 是表格样式
# print("Available table styles:")
# for style in table_styles:
#     print(style)


# ================================================
# ==========    Picture in Table     =============
# ================================================

table_pic = document.add_table(1, 2)
table_pic.style = "Table Grid"
cell = table_pic.cell(0, 0)
paragraph = cell.paragraphs[0]
run = paragraph.add_run()
run.add_picture(pic_square_path, width=Inches(2))
cell = table_pic.cell(0, 1)
paragraph = cell.paragraphs[0]
run = paragraph.add_run()
run.add_picture(pic_square_path, width=Inches(2))

# document.tables

# for paragraph in document.paragraphs:
#     print(paragraph, paragraph.text)


description_str = "this is a test for cp performance."
add_description_demopic_in_docx(insert_index, document, description_str, pic_square_path)



doc_save_path = os.path.join(current_path, 'demo.docx')
document.save(doc_save_path)
